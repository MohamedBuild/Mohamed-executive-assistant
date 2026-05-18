"""
generate-invoice.py
Batch-generates branded AGENCINA PDF invoices from Google Sheet data.

Usage:
  python equipment/generate-invoice.py --batch-file live/invoices/_batch_input.json
  python equipment/generate-invoice.py --setup-template   (creates the placeholder template)
"""

import argparse
import copy
import json
import os
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx2pdf import convert
from dotenv import load_dotenv
from lxml import etree

load_dotenv()

# ── PAID watermark (Word VML / header injection) ──────────────────────────────
# Standard Word watermark: diagonal green "PAID" text rendered via VML TextPath.
# lxml is already installed as a python-docx dependency.
_PAID_WATERMARK_XML = (
    '<w:pict'
    ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    ' xmlns:v="urn:schemas-microsoft-com:vml"'
    ' xmlns:w10="urn:schemas-microsoft-com:office:word">'
    '<v:shape id="_x0000_s2051" type="#_x0000_t136"'
    ' style="position:absolute;margin-left:0;margin-top:0;'
    'width:467.85pt;height:209.55pt;rotation:315;z-index:-251657216;'
    'mso-position-horizontal:center;mso-position-horizontal-relative:margin;'
    'mso-position-vertical:center;mso-position-vertical-relative:margin"'
    ' fillcolor="#ededed" stroked="f">'
    '<v:fill on="t" focussize="0,0"/>'
    '<v:path textpathok="t"/>'
    '<v:textpath on="t" string="PAID"'
    ' style="font-family:Arial;font-size:1pt;font-weight:bold"/>'
    '<w10:wrap type="none"/>'
    '<w10:anchorlock/>'
    '</v:shape>'
    '</w:pict>'
)
_PAID_PICT = etree.fromstring(_PAID_WATERMARK_XML)

TEMPLATE_SRC  = Path("templates/invoice_reference/Agencina_Invoice_INV-2026-0184.docx")
TEMPLATE_PATH = Path("templates/invoice_reference/invoice_template.docx")
OUTPUT_DIR    = Path("live/invoices")


# ── Helpers ───────────────────────────────────────────────────────────────────

def fmt_date(date_str: str) -> str:
    """Convert YYYY-MM-DD to '4 March 2026'."""
    try:
        return datetime.strptime(date_str.strip(), "%Y-%m-%d").strftime("%-d %B %Y")
    except Exception:
        return date_str or ""


def fmt_money(amount, currency: str = "USD") -> str:
    """Format a number as '3,250.00 USD'."""
    try:
        return f"{float(amount):,.2f} {currency}"
    except Exception:
        return f"{amount} {currency}"


def set_run_text(para, text: str):
    """Set the full text of a paragraph using its first run, clearing the rest."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ""


def replace_tokens_in_para(para, replacements: dict):
    """Replace {{TOKEN}} markers across all runs in a paragraph.
    Skips paragraphs with no tokens to avoid wiping embedded images."""
    if not para.runs:
        return
    full = "".join(r.text for r in para.runs)
    if "{{" not in full:
        return
    for token, value in replacements.items():
        full = full.replace(token, str(value))
    para.runs[0].text = full
    for run in para.runs[1:]:
        run.text = ""


def replace_tokens_in_doc(doc: Document, replacements: dict):
    """Walk every paragraph and table cell and replace all tokens."""
    for para in doc.paragraphs:
        replace_tokens_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_tokens_in_para(para, replacements)


# ── Watermark ─────────────────────────────────────────────────────────────────

def add_paid_watermark(doc: Document):
    """Inject a diagonal green PAID watermark into every page via the header."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = para.add_run()
        run._r.append(copy.deepcopy(_PAID_PICT))


# ── Template creation (run once) ──────────────────────────────────────────────

def setup_template():
    """
    Build invoice_template.docx with {{TOKEN}} placeholders.
    Modifies the reference doc by index — no string matching needed.
    """
    if not TEMPLATE_SRC.exists():
        raise FileNotFoundError(f"Reference template not found: {TEMPLATE_SRC}")

    doc = Document(TEMPLATE_SRC)

    # ── Table 0, Col 0: Business info ─────────────────────────────────────────
    c = doc.tables[0].cell(0, 0)
    paras = c.paragraphs
    # [0] blank spacer → replace with logo
    logo_path = Path("assets/Black and White Minimalist Tattoo Studio Logo(1).png")
    if logo_path.exists() and paras:
        for run in paras[0].runs:
            run.text = ""
        paras[0].add_run().add_picture(str(logo_path), width=Cm(3))
    if len(paras) > 1: set_run_text(paras[1], "{{BUSINESS_NAME}}")
    if len(paras) > 2: set_run_text(paras[2], "Agentic Workflow Agency")
    if len(paras) > 3: set_run_text(paras[3], "{{BUSINESS_ADDRESS}}")
    if len(paras) > 4: set_run_text(paras[4], "")
    if len(paras) > 5: set_run_text(paras[5], "{{BUSINESS_EMAIL}}  |  {{BUSINESS_PHONE}}")

    # ── Table 0, Col 1: Invoice number + dates + status ───────────────────────
    c1 = doc.tables[0].cell(0, 1)
    c1p = c1.paragraphs
    # [0]: "INVOICE" — keep
    if len(c1p) > 1:
        set_run_text(c1p[1], "No. {{INVOICE_NUMBER}}")
    c1.add_paragraph("Date:    {{ISSUE_DATE}}")
    c1.add_paragraph("Due:     {{DUE_DATE}}")
    c1.add_paragraph("Status:  {{STATUS}}")

    # ── Table 1, Col 0: Bill To ───────────────────────────────────────────────
    bill = doc.tables[1].cell(0, 0)
    bp = bill.paragraphs
    # [0]: "BILL TO" label — keep
    if len(bp) > 1: set_run_text(bp[1], "{{CLIENT_COMPANY}}")
    for para in bp[2:]:
        set_run_text(para, "")

    # ── Table 2: Line items — keep header row + 1 data row ───────────────────
    t2 = doc.tables[2]
    tbl_xml = t2._tbl
    for row in list(t2.rows[2:]):
        tbl_xml.remove(row._tr)

    data_row = doc.tables[2].rows[1]
    dc = data_row.cells
    if dc[0].paragraphs:
        set_run_text(dc[0].paragraphs[0], "{{DESCRIPTION}}")
        for extra in dc[0].paragraphs[1:]:
            set_run_text(extra, "")
    if dc[1].paragraphs: set_run_text(dc[1].paragraphs[0], "1")
    if dc[2].paragraphs: set_run_text(dc[2].paragraphs[0], "{{SUBTOTAL}} {{CURRENCY}}")
    if dc[3].paragraphs: set_run_text(dc[3].paragraphs[0], "{{SUBTOTAL}} {{CURRENCY}}")

    # ── Table 3: Totals ───────────────────────────────────────────────────────
    t3 = doc.tables[3]
    set_run_text(t3.rows[0].cells[0].paragraphs[0], "Subtotal")
    set_run_text(t3.rows[0].cells[1].paragraphs[0], "{{SUBTOTAL}} {{CURRENCY}}")
    set_run_text(t3.rows[1].cells[0].paragraphs[0], "VAT ({{VAT_PCT}}%)")
    set_run_text(t3.rows[1].cells[1].paragraphs[0], "{{VAT_AMOUNT}} {{CURRENCY}}")
    set_run_text(t3.rows[2].cells[0].paragraphs[0], "TOTAL DUE")
    set_run_text(t3.rows[2].cells[1].paragraphs[0], "{{TOTAL}} {{CURRENCY}}")

    # ── Table 4, Col 0: Payment info ─────────────────────────────────────────
    pay = doc.tables[4].cell(0, 0)
    pay_lines = [
        "PAYMENT INFORMATION",
        "Bank: {{BANK_NAME}}",
        "Account Name: {{BANK_ACCOUNT}}",
        "IBAN: {{BANK_IBAN}}",
        "SWIFT: {{BANK_SWIFT}}",
    ]
    for i, para in enumerate(pay.paragraphs):
        set_run_text(para, pay_lines[i] if i < len(pay_lines) else "")

    # ── Table 4, Col 1: Notes ────────────────────────────────────────────────
    notes = doc.tables[4].cell(0, 1)
    notes_lines = [
        "NOTES & TERMS",
        "Thank you for your business.",
        "Payment is due within 30 days of invoice date.",
        "Please reference {{INVOICE_NUMBER}} with your payment.",
        "",
        "",
    ]
    for i, para in enumerate(notes.paragraphs):
        set_run_text(para, notes_lines[i] if i < len(notes_lines) else "")

    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TEMPLATE_PATH)
    print(f"Template saved: {TEMPLATE_PATH}")
    print("Open it in Word to review placeholders, then run --batch-file to generate PDFs.")


# ── Invoice generation ────────────────────────────────────────────────────────

def build_replacements(inv: dict, env: dict) -> dict:
    """Build the full token → value map for one invoice."""
    currency = inv.get("currency", env.get("CURRENCY", "USD"))

    def to_float(val, default=0.0):
        try:
            return float(str(val).replace(",", ""))
        except (ValueError, TypeError):
            return default

    subtotal   = to_float(inv.get("subtotal", 0))
    vat_pct    = to_float(inv.get("vat_pct", 0))
    total      = to_float(inv.get("total", subtotal))
    vat_amount = round(subtotal * vat_pct / 100, 2)

    status_raw = str(inv.get("status", "")).lower()
    status     = "PAID" if status_raw == "paid" else "OUTSTANDING"

    return {
        # Invoice meta
        "{{INVOICE_NUMBER}}": inv.get("invoice_number", ""),
        "{{ISSUE_DATE}}":     fmt_date(inv.get("issue_date", "")),
        "{{DUE_DATE}}":       fmt_date(inv.get("due_date", "")),
        "{{STATUS}}":         status,
        # Client
        "{{CLIENT_COMPANY}}": inv.get("client_company", ""),
        # Line item
        "{{DESCRIPTION}}":    inv.get("description", ""),
        "{{SUBTOTAL}}":       f"{subtotal:,.2f}",
        "{{VAT_PCT}}":        f"{vat_pct:.0f}" if vat_pct == int(vat_pct) else str(vat_pct),
        "{{VAT_AMOUNT}}":     f"{vat_amount:,.2f}",
        "{{TOTAL}}":          f"{total:,.2f}",
        "{{CURRENCY}}":       currency,
        # Business (from .env)
        "{{BUSINESS_NAME}}":    env.get("BUSINESS_NAME", "AGENCINA"),
        "{{BUSINESS_ADDRESS}}": env.get("BUSINESS_ADDRESS", ""),
        "{{BUSINESS_EMAIL}}":   env.get("BUSINESS_EMAIL", ""),
        "{{BUSINESS_PHONE}}":   env.get("BUSINESS_PHONE", ""),
        # Bank (from .env)
        "{{BANK_NAME}}":    env.get("BANK_NAME", ""),
        "{{BANK_ACCOUNT}}": env.get("BANK_ACCOUNT", ""),
        "{{BANK_IBAN}}":    env.get("BANK_IBAN", ""),
        "{{BANK_SWIFT}}":   env.get("BANK_SWIFT", ""),
    }


def generate_invoice(inv: dict, env: dict) -> Path:
    """Fill template for one invoice and convert to PDF."""
    if not TEMPLATE_PATH.exists():
        print("Template not found — running setup first...")
        setup_template()

    invoice_number = inv.get("invoice_number", "UNKNOWN")
    client_slug    = inv.get("client_slug", "client").replace(" ", "-").lower()
    pdf_name       = f"{invoice_number}_{client_slug}.pdf"
    pdf_path       = OUTPUT_DIR / pdf_name

    doc = Document(TEMPLATE_PATH)
    replacements = build_replacements(inv, env)
    replace_tokens_in_doc(doc, replacements)

    if str(inv.get("status", "")).lower() == "paid":
        add_paid_watermark(doc)

    # Save filled .docx to a temp file, convert to PDF, clean up
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    doc.save(tmp_path)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    convert(tmp_path, str(pdf_path))
    os.unlink(tmp_path)

    return pdf_path


# ── CLI entry point ───────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate AGENCINA PDF invoices.")
    group  = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--setup-template", action="store_true",
                       help="Create invoice_template.docx with {{TOKEN}} placeholders")
    group.add_argument("--batch-file", metavar="JSON_FILE",
                       help="JSON file with array of invoice objects from the sheet")
    args = parser.parse_args()

    if args.setup_template:
        setup_template()
        return

    batch_file = Path(args.batch_file)
    if not batch_file.exists():
        raise FileNotFoundError(f"Batch file not found: {batch_file}")

    with open(batch_file, encoding="utf-8") as fh:
        invoices = json.load(fh)

    env = {
        "BUSINESS_NAME":    os.getenv("BUSINESS_NAME", "AGENCINA"),
        "BUSINESS_ADDRESS": os.getenv("BUSINESS_ADDRESS", ""),
        "BUSINESS_EMAIL":   os.getenv("BUSINESS_EMAIL", ""),
        "BUSINESS_PHONE":   os.getenv("BUSINESS_PHONE", ""),
        "BANK_NAME":        os.getenv("BANK_NAME", ""),
        "BANK_ACCOUNT":     os.getenv("BANK_ACCOUNT", ""),
        "BANK_IBAN":        os.getenv("BANK_IBAN", ""),
        "BANK_SWIFT":       os.getenv("BANK_SWIFT", ""),
        "CURRENCY":         os.getenv("CURRENCY", "USD"),
    }

    print(f"Generating {len(invoices)} invoice(s)...\n")
    generated = []
    for inv in invoices:
        try:
            pdf = generate_invoice(inv, env)
            total = inv.get("total", "?")
            currency = inv.get("currency", "USD")
            print(f"  OK  {pdf.name}  ({total} {currency})")
            generated.append(pdf)
        except Exception as e:
            print(f"  ERR {inv.get('invoice_number', '?')} -- {e}")

    print(f"\n{len(generated)} PDF(s) saved to {OUTPUT_DIR}/")
    print("Nothing has been sent. Review before sharing with clients.")


if __name__ == "__main__":
    main()
